import base64
import random
import re
import sys
import asyncio
import json
import os
from pathlib import Path

import pandas as pd
import pdfplumber
import requests
import streamlit as st
from docx import Document


BASE_DIR = Path(__file__).resolve().parent
DOCS_DIR = BASE_DIR / "docs"

SUBJECTS = {
    "oop": {
        "label": "OOP (Python)",
        "files": ["OOP_30_AIKEN_EN_v2.txt", "oop_extra_tasks_en.docx"],
    },
    "nm": {"label": "Numerical Methods", "files": ["NM.pdf"]},
    "diff": {"label": "Differential Equations", "files": ["DIff.pdf"]},
}

FILL_SYNONYMS = {
    "objects": {"object", "objects", "instances"},
    "constructor": {"constructor", "init", "__init__", "__init__ method"},
    "the current object": {"the current object", "current object", "self", "instance", "the instance"},
    "private": {"private", "private attribute", "encapsulated", "hidden"},
    "super()": {"super()", "super", "super function", "super() function"},
}

GROQ_VISION_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"


def clean_text(text):
    return re.sub(r"\s+", " ", text or "").strip()


def normalize_answer(value):
    value = str(value or "").strip().lower()
    value = value.replace("—", "-").replace("–", "-")
    value = re.sub(r"[_\s]+", " ", value)
    value = re.sub(r"[^a-z0-9_()+.\- /]", "", value)
    return value.strip()


def make_question(q, opts=None, ans="", qtype="mcq", solution=""):
    return {
        "q": clean_text(q),
        "opts": [clean_text(o) for o in opts] if opts else None,
        "ans": clean_text(ans),
        "type": qtype,
        "solution": clean_text(solution),
    }


def answer_letter_index(answer, opts=None):
    answer = clean_text(answer).upper()
    if len(answer) != 1 or not answer.isalpha():
        return None
    idx = ord(answer) - ord("A")
    if opts is not None and not (0 <= idx < len(opts)):
        return None
    return idx


def parse_aiken(path):
    text = path.read_text(encoding="utf-8-sig")
    questions = []
    blocks = re.split(r"\n\s*\n", text.strip())
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        answer_line = next((line for line in lines if line.upper().startswith("ANSWER:")), "")
        if not answer_line:
            continue

        answer = answer_line.split(":", 1)[1].strip().upper()[:1]
        option_lines = [line for line in lines if re.match(r"^[A-D]\.\s+", line)]
        if len(option_lines) < 2 or answer not in "ABCD":
            continue

        first_option = lines.index(option_lines[0])
        q_text = " ".join(lines[:first_option])
        opts = [re.sub(r"^[A-D]\.\s*", "", line) for line in option_lines]
        questions.append(make_question(q_text, opts, answer, "mcq"))
    return questions


def paragraph_text(paragraph):
    return clean_text(paragraph.text)


def parse_extra_docx(path):
    doc = Document(path)
    lines = [paragraph_text(p) for p in doc.paragraphs]
    lines = [line for line in lines if line]

    fill_questions = {}
    fill_answers = {}
    tf_questions = {}
    tf_answers = {}
    section = None

    for line in lines:
        lower = line.lower()
        if lower.startswith("answer key") and "part ii" in lower:
            section = "tf_key"
            continue
        if lower.startswith("answer key") and "part i" in lower:
            section = "fill_key"
            continue
        if lower.startswith("part ii"):
            section = "tf"
            continue
        if lower.startswith("part i"):
            section = "fill"
            continue
        if lower.startswith("instructions") or lower.startswith("answer:"):
            continue

        m = re.match(r"^(\d+)\.\s*(.+)$", line)
        if not m:
            continue
        number = int(m.group(1))
        body = clean_text(m.group(2))

        if section == "fill":
            fill_questions[number] = body
        elif section == "fill_key":
            fill_answers[number] = body
        elif section == "tf":
            tf_questions[number] = body
        elif section == "tf_key":
            parts = re.split(r"\s+[-—]\s+", body, maxsplit=1)
            tf_answers[number] = (parts[0].strip(), parts[1].strip() if len(parts) > 1 else "")

    questions = []
    for number in sorted(fill_questions):
        answer = fill_answers.get(number, "")
        questions.append(make_question(fill_questions[number], None, answer, "fill"))

    for number in sorted(tf_questions):
        answer, explanation = tf_answers.get(number, ("", ""))
        if answer:
            questions.append(make_question(tf_questions[number], ["True", "False"], answer, "mcq", explanation))

    return questions


def extract_pdf_text(path):
    chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            chunks.append(page.extract_text() or "")
    return "\n".join(chunks)


def parse_pdf_mcq(path):
    text = extract_pdf_text(path)
    if not clean_text(text):
        return []

    text = text.replace("\r", "\n")
    text = re.sub(r"(?m)^\s*Question\s+\d+.*$", "", text)
    text = re.sub(r"(?m)^\s*(Not yet answered|Marked out of 1|Flag question|Select one:).*$", "", text)

    answer_pattern = re.compile(r"(?:ANSWER|Answer|Correct answer)\s*[:\-]\s*([A-D])", re.I)
    starts = [m.start() for m in re.finditer(r"(?m)^\s*(?:\d+[\).]\s+|[Qq]uestion\s+\d+)", text)]
    if not starts:
        starts = [0]
    starts.append(len(text))

    questions = []
    for start, end in zip(starts, starts[1:]):
        block = text[start:end].strip()
        if not block:
            continue
        answer_match = answer_pattern.search(block)
        if not answer_match:
            continue
        answer = answer_match.group(1).upper()
        block = answer_pattern.sub("", block)

        option_matches = list(re.finditer(r"(?ms)^\s*([A-D])[\).]\s*(.+?)(?=^\s*[A-D][\).]\s*|\Z)", block))
        if len(option_matches) < 2:
            continue
        first_option_start = option_matches[0].start()
        q_text = re.sub(r"^\s*(?:\d+[\).]\s+|[Qq]uestion\s+\d+\s*)", "", block[:first_option_start]).strip()
        opts = [m.group(2).strip() for m in option_matches]
        questions.append(make_question(q_text, opts, answer, "mcq"))

    return questions


def load_generated_questions(path):
    if not path.exists():
        return []
    data = json.loads(path.read_text(encoding="utf-8"))
    questions = data.get("questions", data if isinstance(data, list) else [])
    parsed = []
    for item in questions:
        q = clean_text(item.get("q", ""))
        opts = item.get("opts")
        ans = clean_text(item.get("ans", "")).upper()[:1]
        if (
            isinstance(opts, list)
            and ans == "E"
            and len(opts) == 4
            and "all" in normalize_answer(item.get("solution", ""))
        ):
            opts = opts + ["All of the above"]
        if q and isinstance(opts, list) and len(opts) >= 2 and answer_letter_index(ans, opts) is not None:
            parsed.append(make_question(q, opts, ans, "mcq", item.get("solution", "")))
    return parsed


def parse_ocr_lines_to_question(lines):
    text_lines = [clean_text(line) for line in lines if clean_text(line)]
    if not text_lines:
        return None

    body = []
    collecting = False
    for line in text_lines:
        lower = line.lower()
        if lower.startswith("which ") or lower.startswith("what ") or lower.startswith("find ") or lower.startswith("solve "):
            collecting = True
        if collecting:
            if lower.startswith("next page") or lower.startswith("quiz navigation") or lower.startswith("finish attempt"):
                break
            body.append(line)

    if not body:
        return None

    joined = " ".join(body)
    joined = re.sub(r"\bSelect one:\b", " Select one: ", joined, flags=re.I)
    question_part = joined.split("Select one:", 1)[0].strip()
    option_part = joined.split("Select one:", 1)[1].strip() if "Select one:" in joined else ""

    option_matches = list(re.finditer(r"\b([A-D])[\).]\s*", option_part, re.I))
    opts = []
    if len(option_matches) >= 2:
        for idx, match in enumerate(option_matches):
            start = match.end()
            end = option_matches[idx + 1].start() if idx + 1 < len(option_matches) else len(option_part)
            opt = clean_text(option_part[start:end])
            if opt:
                opts.append(opt)

    if len(opts) < 2:
        return None

    return make_question(question_part, opts, "", "mcq", "OCR extracted from a scanned PDF. Add an answer key to score this item.")


def render_pdf_page_image(pdf_path, page_index, output_dir, scale=4):
    import pypdfium2 as pdfium

    output_dir.mkdir(parents=True, exist_ok=True)
    image_path = output_dir / f"{pdf_path.stem}_page_{page_index + 1:02d}.png"
    pdf = pdfium.PdfDocument(str(pdf_path))
    page = pdf[page_index]
    page.render(scale=scale).to_pil().save(image_path)
    return image_path


def render_pdf_page_jpeg(pdf_path, page_index, output_dir, scale=2.4, quality=78):
    import pypdfium2 as pdfium

    output_dir.mkdir(parents=True, exist_ok=True)
    image_path = output_dir / f"{pdf_path.stem}_page_{page_index + 1:02d}.jpg"
    pdf = pdfium.PdfDocument(str(pdf_path))
    page = pdf[page_index]
    image = page.render(scale=scale).to_pil().convert("RGB")
    image.save(image_path, format="JPEG", quality=quality, optimize=True)
    return image_path


async def windows_ocr_image(image_path):
    from winrt.windows.globalization import Language
    from winrt.windows.graphics.imaging import BitmapDecoder
    from winrt.windows.media.ocr import OcrEngine
    from winrt.windows.storage import FileAccessMode, StorageFile

    storage_file = await StorageFile.get_file_from_path_async(str(image_path.resolve()))
    stream = await storage_file.open_async(FileAccessMode.READ)
    decoder = await BitmapDecoder.create_async(stream)
    bitmap = await decoder.get_software_bitmap_async()
    engine = OcrEngine.try_create_from_language(Language("en-US"))
    result = await engine.recognize_async(bitmap)
    return [line.text for line in result.lines]


async def ocr_pdf_to_text(pdf_path, output_txt, images_dir):
    import pypdfium2 as pdfium

    pages = len(pdfium.PdfDocument(str(pdf_path)))
    chunks = []
    for page_index in range(pages):
        image_path = render_pdf_page_image(pdf_path, page_index, images_dir)
        lines = await windows_ocr_image(image_path)
        chunks.append(f"--- Page {page_index + 1} ---\n" + "\n".join(lines))
    output_txt.write_text("\n\n".join(chunks), encoding="utf-8")
    return pages


def convert_scanned_pdfs_with_ocr():
    targets = [DOCS_DIR / "NM.pdf", DOCS_DIR / "DIff.pdf"]
    output_dir = DOCS_DIR / "ocr_images"
    for pdf_path in targets:
        output_txt = DOCS_DIR / f"{pdf_path.stem}_ocr.txt"
        pages = asyncio.run(ocr_pdf_to_text(pdf_path, output_txt, output_dir))
        print(f"{pdf_path.name}: OCR converted {pages} pages -> {output_txt.relative_to(BASE_DIR)}")
    print("Note: OCR text does not include correct answers unless they are visible in the source scans.")


def groq_extract_image_question(image_path, subject, page_number):
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        raise RuntimeError("Set GROQ_API_KEY before running --groq-extract-pdfs.")

    encoded = base64.b64encode(image_path.read_bytes()).decode("ascii")
    prompt = f"""
Extract the quiz question from this scanned {subject} screenshot.

Return only valid JSON with exactly this shape:
{{
  "q": "question text",
  "opts": ["option A text", "option B text", "option C text", "option D text"],
  "ans": "A|B|C|D",
  "solution": "brief explanation of why the answer is correct"
}}

Rules:
- Preserve mathematical notation as readable plain text.
- Use ASCII-only text for math symbols: write integral as integral, arrows as ->, infinity as infinity.
- Ignore browser UI, timers, page navigation, and radio button circles.
- If an option letter is visually confused by OCR, infer the A/B/C/D order from top to bottom.
- The screenshot does not show the correct option. Solve the problem yourself and set ans to the correct letter.
- If the page is unreadable, return {{"q":"","opts":[],"ans":"","solution":"unreadable page"}}.
- This is page {page_number}.
""".strip()

    payload = {
        "model": GROQ_VISION_MODEL,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/jpeg;base64,{encoded}"},
                    },
                ],
            }
        ],
        "temperature": 0,
        "max_completion_tokens": 1200,
        "response_format": {"type": "json_object"},
    }
    response = requests.post(
        GROQ_API_URL,
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json=payload,
        timeout=90,
    )
    response.raise_for_status()
    content = response.json()["choices"][0]["message"]["content"]
    return json.loads(content)


def groq_extract_pdf_questions(pdf_path, subject, max_pages=None):
    import pypdfium2 as pdfium

    pages = len(pdfium.PdfDocument(str(pdf_path)))
    if max_pages:
        pages = min(pages, max_pages)

    image_dir = DOCS_DIR / "groq_images"
    questions = []
    text_chunks = []
    for page_index in range(pages):
        page_number = page_index + 1
        image_path = render_pdf_page_jpeg(pdf_path, page_index, image_dir)
        print(f"{pdf_path.name}: extracting page {page_number}/{pages}")
        item = groq_extract_image_question(image_path, subject, page_number)
        item["source_page"] = page_number
        text_chunks.append(f"--- Page {page_number} ---\n{json.dumps(item, ensure_ascii=False, indent=2)}")
        if item.get("q") and item.get("opts") and item.get("ans"):
            questions.append(item)

    output_json = DOCS_DIR / f"{pdf_path.stem}_groq_questions.json"
    output_txt = DOCS_DIR / f"{pdf_path.stem}_groq_ocr.txt"
    output_json.write_text(json.dumps({"questions": questions}, ensure_ascii=False, indent=2), encoding="utf-8")
    output_txt.write_text("\n\n".join(text_chunks), encoding="utf-8")
    print(f"{pdf_path.name}: saved {len(questions)} questions -> {output_json.relative_to(BASE_DIR)}")
    print(f"{pdf_path.name}: saved extraction text -> {output_txt.relative_to(BASE_DIR)}")


def groq_extract_scanned_pdfs():
    max_pages = None
    if "--max-pages" in sys.argv:
        idx = sys.argv.index("--max-pages")
        if idx + 1 < len(sys.argv):
            max_pages = int(sys.argv[idx + 1])

    targets = [
        (DOCS_DIR / "NM.pdf", "Numerical Methods"),
        (DOCS_DIR / "DIff.pdf", "Differential Equations"),
    ]
    for pdf_path, subject in targets:
        groq_extract_pdf_questions(pdf_path, subject, max_pages=max_pages)


def load_all_questions():
    oop_aiken = parse_aiken(DOCS_DIR / "OOP_30_AIKEN_EN_v2.txt")
    oop_extra = parse_extra_docx(DOCS_DIR / "oop_extra_tasks_en.docx")
    nm = parse_pdf_mcq(DOCS_DIR / "NM.pdf")
    diff = parse_pdf_mcq(DOCS_DIR / "DIff.pdf")
    if not nm:
        nm = load_generated_questions(DOCS_DIR / "NM_groq_questions.json")
    if not diff:
        diff = load_generated_questions(DOCS_DIR / "DIff_groq_questions.json")
    return {
        "oop_aiken": oop_aiken,
        "oop_extra": oop_extra,
        "oop": oop_aiken + oop_extra,
        "nm": nm,
        "diff": diff,
    }


def answer_is_correct(question, user_answer):
    correct = question["ans"]
    if question["type"] == "mcq":
        if answer_letter_index(correct, question.get("opts") or []) is not None:
            return str(user_answer).upper() == correct.upper()
        idx = answer_letter_index(str(user_answer), question.get("opts") or [])
        if idx is not None:
            user_answer = question["opts"][idx]
        return normalize_answer(user_answer) == normalize_answer(correct)

    user_norm = normalize_answer(user_answer)
    correct_norm = normalize_answer(correct)
    accepted = {correct_norm}
    accepted.update(normalize_answer(item) for item in FILL_SYNONYMS.get(correct_norm, set()))
    return user_norm in accepted


def correct_answer_text(question):
    answer = question["ans"]
    if question["type"] == "mcq" and question["opts"]:
        idx = answer_letter_index(answer, question["opts"])
        if idx is not None:
            return f"{answer.upper()}. {question['opts'][idx]}"
    return answer


def start_quiz(subject_key):
    questions = list(st.session_state.question_bank.get(subject_key, []))
    random.shuffle(questions)
    st.session_state.subject = subject_key
    st.session_state.questions = questions
    st.session_state.answers = {}
    st.session_state.index = 0
    st.session_state.show_review = False
    st.session_state.page = "quiz"


def retry_quiz():
    start_quiz(st.session_state.subject)


def go_home():
    st.session_state.page = "home"
    st.session_state.subject = None
    st.session_state.questions = []
    st.session_state.answers = {}
    st.session_state.index = 0
    st.session_state.show_review = False


def sidebar_progress():
    questions = st.session_state.get("questions", [])
    if not questions:
        return
    total = len(questions)
    answered = len(st.session_state.get("answers", {}))
    st.sidebar.header(SUBJECTS[st.session_state.subject]["label"])
    st.sidebar.progress(answered / total)
    st.sidebar.write(f"Question {min(st.session_state.index + 1, total)} of {total}")
    st.sidebar.write(f"Answered: {answered}/{total}")
    if st.sidebar.button("Home"):
        go_home()
        st.rerun()


def render_home():
    st.title("Self-Test Quiz")
    st.caption("Choose a subject. Questions are shuffled every attempt.")

    counts = {
        "OOP (Python)": len(st.session_state.question_bank["oop"]),
        "Numerical Methods": len(st.session_state.question_bank["nm"]),
        "Differential Equations": len(st.session_state.question_bank["diff"]),
    }
    st.dataframe(pd.DataFrame(counts.items(), columns=["Subject", "Parsed questions"]), hide_index=True)

    cols = st.columns(3)
    for col, subject_key in zip(cols, ["oop", "nm", "diff"]):
        with col:
            disabled = len(st.session_state.question_bank[subject_key]) == 0
            if st.button(SUBJECTS[subject_key]["label"], use_container_width=True, disabled=disabled):
                start_quiz(subject_key)
                st.rerun()

    if not st.session_state.question_bank["nm"] or not st.session_state.question_bank["diff"]:
        st.warning(
            "NM.pdf and DIff.pdf are scanned/image PDFs. pdfplumber found no selectable text, "
            "so those subjects cannot be scored until OCR text or text-based PDFs with answers are provided."
        )


def render_quiz():
    sidebar_progress()
    questions = st.session_state.questions
    if not questions:
        st.warning("No parsed questions are available for this subject.")
        if st.button("Home"):
            go_home()
            st.rerun()
        return

    idx = st.session_state.index
    question = questions[idx]
    st.subheader(f"Question {idx + 1} of {len(questions)}")
    st.write(question["q"])

    current_answer = st.session_state.answers.get(idx, "")
    if question["type"] == "mcq":
        labels = [f"{chr(65 + i)}. {opt}" for i, opt in enumerate(question["opts"] or [])]
        current_index = None
        current_index = answer_letter_index(current_answer, question["opts"] or [])
        if current_index is None and question["opts"]:
            for option_index, option in enumerate(question["opts"]):
                if normalize_answer(current_answer) == normalize_answer(option):
                    current_index = option_index
                    break
        selected = st.radio("Choose one answer", labels, index=current_index, key=f"q_{idx}")
        if selected and answer_letter_index(question["ans"], question["opts"] or []) is not None:
            pending_answer = selected.split(".", 1)[0]
        elif selected:
            pending_answer = selected.split(".", 1)[1].strip()
        else:
            pending_answer = ""
    else:
        pending_answer = st.text_input("Your answer", value=current_answer, key=f"q_{idx}")

    nav = st.columns([1, 1, 2])
    with nav[0]:
        if st.button("Previous", disabled=idx == 0):
            st.session_state.answers[idx] = pending_answer
            st.session_state.index -= 1
            st.rerun()
    with nav[1]:
        is_last = idx == len(questions) - 1
        label = "Finish" if is_last else "Next"
        if st.button(label, type="primary"):
            st.session_state.answers[idx] = pending_answer
            if is_last:
                st.session_state.page = "results"
            else:
                st.session_state.index += 1
            st.rerun()


def render_results():
    sidebar_progress()
    questions = st.session_state.questions
    answers = st.session_state.answers
    total = len(questions)
    correct = sum(answer_is_correct(q, answers.get(i, "")) for i, q in enumerate(questions))
    percent = round((correct / total) * 100, 1) if total else 0

    st.title("Results")
    st.metric("Score", f"{percent}%", f"{correct}/{total}")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Retry", use_container_width=True):
            retry_quiz()
            st.rerun()
    with col2:
        if st.button("Show Full Review", use_container_width=True):
            st.session_state.show_review = not st.session_state.show_review

    if st.session_state.show_review:
        st.subheader("Full Review")
        for i, question in enumerate(questions):
            user_answer = answers.get(i, "")
            ok = answer_is_correct(question, user_answer)
            with st.expander(f"Question {i + 1}: {'Correct' if ok else 'Incorrect'}"):
                st.write(question["q"])
                color = "#16794c" if ok else "#b42318"
                st.markdown(f"<div style='color:{color}; font-weight:700'>Your answer: {user_answer or 'Blank'}</div>", unsafe_allow_html=True)
                st.markdown(f"<div style='color:#16794c; font-weight:700'>Correct answer: {correct_answer_text(question)}</div>", unsafe_allow_html=True)
                solution = question.get("solution") or "Review the source material for this item."
                st.info(solution)


def print_summary():
    bank = load_all_questions()
    print("Parsed question summary:")
    print(f"docs/OOP_30_AIKEN_EN_v2.txt: {len(bank['oop_aiken'])}")
    print(f"docs/oop_extra_tasks_en.docx: {len(bank['oop_extra'])}")
    print(f"docs/NM.pdf: {len(bank['nm'])}")
    print(f"docs/DIff.pdf: {len(bank['diff'])}")
    print(f"OOP total: {len(bank['oop'])}")


def main():
    st.set_page_config(page_title="Self-Test Quiz", page_icon="?", layout="centered")
    if "question_bank" not in st.session_state:
        st.session_state.question_bank = load_all_questions()
    st.session_state.setdefault("page", "home")
    st.session_state.setdefault("subject", None)
    st.session_state.setdefault("questions", [])
    st.session_state.setdefault("answers", {})
    st.session_state.setdefault("index", 0)
    st.session_state.setdefault("show_review", False)

    if st.session_state.page == "home":
        render_home()
    elif st.session_state.page == "quiz":
        render_quiz()
    elif st.session_state.page == "results":
        render_results()


if __name__ == "__main__":
    if "--summary" in sys.argv:
        print_summary()
    elif "--ocr-pdfs" in sys.argv:
        convert_scanned_pdfs_with_ocr()
    elif "--groq-extract-pdfs" in sys.argv:
        groq_extract_scanned_pdfs()
    else:
        main()
